# 由 Cline 协助修改
import os
import re
import time

import streamlit as st
import requests
from datetime import datetime
from docx import Document
from docx.shared import Pt
from io import BytesIO
from pathlib import Path

API_URL = "https://api.siliconflow.cn/v1/chat/completions"
MAX_RETRIES = 3
RETRY_DELAY = 2
MODEL_OPTIONS = {
    "DeepSeek-V3": "deepseek-ai/DeepSeek-V3", 
    "Kimi-K2.5": "Pro/moonshotai/Kimi-K2.5",
    "通义千问 Qwen2.5-72B": "Qwen/Qwen2.5-72B-Instruct",
    "通义千问 Qwen3-8B": "Qwen/Qwen3-8B",
    }
REQUEST_TIMEOUT = 60
LOG_DIR = Path(__file__).parent
MAX_REWRITE_ATTEMPTS = 3


def call_api(api_key, model_id, prompt, max_tokens=1000):
    """通用API调用函数"""
    headers = {
        "Authorization": f"Bearer {api_key.strip()}",
        "Content-Type": "application/json",
    }
    data = {
        "model": model_id,
        "messages": [{"role": "user", "content": prompt}],
        "max_tokens": max_tokens,
    }
    
    for attempt in range(MAX_RETRIES):
        try:
            response = requests.post(
                API_URL,
                headers=headers,
                json=data,
                timeout=REQUEST_TIMEOUT,
            )
            if response.status_code == 200:
                result = response.json()
                return result["choices"][0]["message"]["content"]
            elif response.status_code in [500, 502, 503, 504]:
                time.sleep(RETRY_DELAY * (attempt + 1))
                continue
            else:
                return f"❌ API错误：{response.status_code}"
        except requests.RequestException as e:
            if attempt < MAX_RETRIES - 1:
                time.sleep(RETRY_DELAY * (attempt + 1))
                continue
            return f"❌ 网络错误：{e}"


def get_plan_prompt(style, product):
    """规划阶段提示词"""
    return f"""你是一个专业的文案策划师。请为以下商品规划文案生成步骤。

商品：{product}
风格：{style}

请输出：
1. 核心卖点（3-4个）
2. 文案结构规划
3. 目标受众分析

请以清晰的列表形式输出，不要超过100字。"""


def get_eval_prompt(style, product, content):
    """评估阶段提示词"""
    return f"""你是一个专业的文案评估师。请评估以下文案的质量。

商品：{product}
风格：{style}
文案内容：
{content}

评估标准：
1. 风格匹配度（0-5分）：是否符合指定风格
2. 说服力（0-5分）：是否能打动目标用户
3. 完整性（0-5分）：是否包含所有必要信息

请严格按照以下JSON格式输出：
{{
    "score": 总分,
    "pass": true或false（总分>=10为pass）,
    "feedback": "具体改进建议"
}}

只输出JSON，不要其他内容。"""


def get_rewrite_prompt(style, product, content, feedback):
    """重写阶段提示词"""
    return f"""请根据以下反馈重写文案。

商品：{product}
风格：{style}
原文案：
{content}

改进建议：
{feedback}

请根据建议进行优化重写，保持风格一致。直接输出新文案。"""


# 风格对应的 prompt 映射
# 页面配置
st.set_page_config(page_title="AI 文案生成器", page_icon="✍️")

# 标题
st.title("✍️ AI 文案生成器")
st.caption(f"当前时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")


# 风格对应的 prompt 映射
def get_prompt(style, product):
    prompts = {
        "小红书种草风": f"""你是一个小红书种草博主，请为以下商品写一篇小红书风格的种草文案。

商品：{product}

要求：
1. 标题要吸引人，带emoji
2. 正文分3-4个卖点，每个卖点用emoji开头
3. 语气亲切自然，像朋友推荐
4. 结尾加上3-5个相关标签

请直接输出文案，不要有其他说明。""",

        "朋友圈分享风": f"""请为商品「{product}」写一段微信朋友圈风格的推荐文案。

要求：
1. 简短有力，不超过150字
2. 语气亲切，像朋友在分享
3. 带2-3个emoji
4. 结尾加一句号召性的话

请直接输出文案。""",

        "专业评测风": f"""请为商品「{product}」写一篇专业评测风格的文案。

要求：
1. 标题客观专业
2. 分3-4个技术/功能卖点，每个卖点有具体说明
3. 语气中立可信
4. 结尾给出总结和购买建议

请直接输出文案。""",

        "淘宝促销风": f"""请为商品「{product}」写一段淘宝详情页风格的促销文案。

要求：
1. 突出卖点和优惠
2. 使用短句、符号分隔
3. 语气热情有感染力
4. 结尾加上行动呼吁（如"点击购买"）

请直接输出文案。""",

        "微博热搜风": f"""请为商品「{product}」写一段微博热搜风格的推荐文案。

要求：
1. 开头用"#话题#"格式，像是热搜话题
2. 正文简短有力，不超过100字
3. 语气像网友热议，带2-3个emoji
4. 结尾带2-3个相关话题标签

请直接输出文案。""",

        "抖音脚本风": f"""请为商品「{product}」写一个抖音短视频脚本风格的文案。

要求：
1. 开头写出"【画面】"和"【文案】"
2. 分3-4个镜头，每个镜头15字以内
3. 最后一句是"🔥 点击购物车，同款 Get！"
4. 整体节奏快，有冲击力

请直接输出文案。""",

        "知乎干货风": f"""请为商品「{product}」写一段知乎回答风格的推荐文案。

要求：
1. 开头用"谢邀"或"作为XXX用户"
2. 正文分3-4个要点，每个要点有具体说明
3. 语气理性、专业、可信
4. 结尾总结并给出建议

请直接输出文案。""",

        "英文国际风": f"""Please write an English product description for "{product}".

Requirements:
1. Short and catchy title
2. 3-4 bullet points highlighting key features
3. Friendly, conversational tone
4. End with a call to action

Output in English only."""
    }
    return prompts.get(style, prompts["小红书种草风"])


def append_copy_log(log_filename, prod, style, model_name, body):
    try:
        log_path = LOG_DIR / log_filename
        now_str = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        with open(log_path, "a", encoding="utf-8") as f:
            f.write(f"\n📝 商品: {prod}\n")
            f.write(f"🎨 风格: {style}\n")
            f.write(f"🤖 模型: {model_name}\n")
            f.write(f"📝 生成时间: {now_str}\n")
            f.write(body + "\n")
            f.write("-" * 50 + "\n")
    except OSError as e:
        st.warning(f"日志写入失败（{prod}）：{e}")


def parse_api_error(response):
    err_msg = f"❌ 生成失败，状态码：{response.status_code}"
    detail = ""
    try:
        body = response.json()
        if isinstance(body.get("error"), dict):
            detail = body["error"].get("message", "")
        if not detail:
            detail = body.get("message", "")
    except ValueError:
        detail = response.text[:300].strip()
    if detail:
        err_msg += f"\n详情：{detail}"
    if response.status_code == 403:
        err_msg += (
            "\n\n💡 Kimi 等部分模型需 SiliconFlow 账户完成实名认证，"
            "且 API Key 须来自已认证账户。请登录 https://cloud.siliconflow.cn "
            "检查「用户中心 → 实名认证」与账户余额。"
        )
    return err_msg


# 侧边栏：API Key 设置
with st.sidebar:
    st.header("⚙️ 设置")
    # 优先 secrets → 环境变量 → 手动输入
    api_key = st.secrets.get("API_KEY", "")
    if api_key:
        st.success("✅ API Key 已从 secrets 自动加载")
    elif os.environ.get("API_KEY"):
        api_key = os.environ.get("API_KEY", "")
        st.success("✅ API Key 已从环境变量自动加载")
    else:
        api_key = st.text_input("API Key", type="password", value="", help="本地测试时手动输入")
    st.markdown("---")
    st.caption("💡 提示：文案会保存在日志文件中，也可导出 Word 文档")

# 主界面
st.header("📝 生成文案")

# 商品名称输入（支持多个，用逗号分隔）
product = st.text_area(
    "商品名称（多个请用逗号分隔）",
    placeholder="例如：无线蓝牙耳机,便携咖啡杯,智能手表",
    height=80,
)

# 风格选择（8种风格）
style = st.selectbox(
    "选择文案风格", 
    options=["小红书种草风", "朋友圈分享风", "专业评测风", "淘宝促销风", "微博热搜风", "抖音脚本风", "知乎干货风", "英文国际风"],
    index=0,
    key="style_select"  # 添加唯一 key
)

# 模型选择
model_label = st.selectbox(
    "选择 AI 模型",
    options=list(MODEL_OPTIONS.keys()),
    index=0,
)
model_id = MODEL_OPTIONS[model_label]

# 初始化 session_state
if "all_results" not in st.session_state:
    st.session_state.all_results = None
if "ok_count" not in st.session_state:
    st.session_state.ok_count = 0
if "fail_count" not in st.session_state:
    st.session_state.fail_count = 0

# 生成按钮
if st.button("🚀 生成文案", type="primary"):
    if not product or not product.strip():
        st.error("请输入商品名称")
    elif not api_key.strip():
        st.error("请在侧边栏填写 API Key")
    else:
        products = [p.strip() for p in product.replace("，", ",").split(",") if p.strip()]
        if not products:
            st.error("请输入至少一个有效商品名称")
        else:
            batch_id = datetime.now().strftime("%Y%m%d_%H%M%S")
            all_results = []
            agent_info = []
            ok_count = 0
            fail_count = 0
            log_filename = f"copy_log_web_{datetime.now().strftime('%Y%m%d')}.log"

            for prod in products:
                with st.status(f"🤖 正在处理「{prod}」...", expanded=True) as status:
                    # 阶段1：规划
                    st.write("🧠 规划中...")
                    plan = call_api(api_key, model_id, get_plan_prompt(style, prod), max_tokens=300)
                    st.write(f"📋 规划完成：{plan[:50]}...")

                    # 阶段2：生成
                    st.write("✍️ 生成中...")
                    prompt = get_prompt(style, prod)
                    content = call_api(api_key, model_id, prompt, max_tokens=1000)
                    if content.startswith("❌"):
                        status.update(label=f"❌ 「{prod}」生成失败", state="error")
                        all_results.append((prod, content))
                        agent_info.append({"plan": plan, "score": 0, "rewrite_count": 0, "feedback": ""})
                        append_copy_log(log_filename, prod, style, model_label, content)
                        fail_count += 1
                        continue

                    # 阶段3：评估与重写循环
                    rewrite_count = 0
                    eval_score = 0
                    feedback = ""
                    passed = False

                    while rewrite_count < MAX_REWRITE_ATTEMPTS and not passed:
                        st.write("🔍 评估中...")
                        eval_result = call_api(api_key, model_id, get_eval_prompt(style, prod, content), max_tokens=300)
                        
                        try:
                            eval_data = eval_result.replace("```json", "").replace("```", "").strip()
                            eval_json = eval_data
                            import json
                            eval_obj = json.loads(eval_json)
                            eval_score = eval_obj.get("score", 0)
                            passed = eval_obj.get("pass", False)
                            feedback = eval_obj.get("feedback", "")
                        except:
                            passed = True

                        if passed:
                            st.write(f"✅ 评估通过（分数：{eval_score}）")
                            break

                        rewrite_count += 1
                        st.write(f"🔄 重写中（第 {rewrite_count} 次）...")
                        content = call_api(api_key, model_id, get_rewrite_prompt(style, prod, content, feedback), max_tokens=1000)
                        if content.startswith("❌"):
                            break

                    if rewrite_count > 0:
                        st.write(f"📝 共重写 {rewrite_count} 次")
                    status.update(label=f"✅ 「{prod}」完成", state="complete")

                    all_results.append((prod, content))
                    agent_info.append({
                        "plan": plan,
                        "score": eval_score,
                        "rewrite_count": rewrite_count,
                        "feedback": feedback
                    })
                    append_copy_log(log_filename, prod, style, model_label, content)
                    ok_count += 1

            # 保存到 session_state
            st.session_state.all_results = all_results
            st.session_state.agent_info = agent_info
            st.session_state.ok_count = ok_count
            st.session_state.fail_count = fail_count
            st.session_state.batch_id = batch_id
            st.session_state.model_label = model_label
            st.session_state.style = style

            if ok_count and not fail_count:
                st.success(f"✅ 成功生成 {ok_count} 篇文案！")
            elif ok_count and fail_count:
                st.warning(f"完成：成功 {ok_count} 篇，失败 {fail_count} 篇")
            else:
                st.error(f"全部失败（{fail_count} 篇），请检查 API Key 或网络")

# 显示结果（从 session_state 读取）
if st.session_state.all_results:
    all_results = st.session_state.all_results
    ok_count = st.session_state.ok_count
    fail_count = st.session_state.fail_count
    batch_id = st.session_state.batch_id
    model_label = st.session_state.model_label
    style = st.session_state.style

    doc = Document()
    doc.add_heading("AI 文案生成报告", 0)
    doc.add_paragraph(f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph(f"AI 模型：{model_label}")
    doc.add_paragraph(f"文案风格：{style}")
    doc.add_paragraph(f"商品数量：{len(all_results)} 个（成功 {ok_count}，失败 {fail_count}）")
    doc.add_paragraph("-" * 50)

    for i, (prod, content) in enumerate(all_results):
        doc.add_heading(f"商品：{prod}", level=1)
        for line in content.split('\n'):
            line = line.strip()
            if not line:
                continue
            
            # 二级标题
            if line.startswith('## '):
                doc.add_heading(line[3:].strip(), level=2)
            # 三级标题
            elif line.startswith('### '):
                doc.add_heading(line[4:].strip(), level=3)
            else:
                # 处理加粗
                p = doc.add_paragraph()
                parts = re.split(r'(\*\*.*?\*\*)', line)
                for part in parts:
                    if part.startswith('**') and part.endswith('**'):
                        run = p.add_run(part[2:-2])
                        run.bold = True
                        run.font.size = Pt(14)  # 加大字体
                    elif part:
                        run = p.add_run(part)
                        run.font.size = Pt(12)  # 正常字体
        
        if i < len(all_results) - 1:
            doc.add_page_break()

    doc_bytes = BytesIO()
    doc.save(doc_bytes)
    doc_bytes.seek(0)

    st.download_button(
        label="📥 下载 Word 文档",
        data=doc_bytes,
        file_name=f"文案报告_{batch_id}.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        key="download_docx",
    )
    st.markdown("---")

    for idx, (prod, content) in enumerate(all_results):
        with st.expander(f"📦 {prod}", expanded=True):
            if "agent_info" in st.session_state and idx < len(st.session_state.agent_info):
                info = st.session_state.agent_info[idx]
                st.markdown(f"**📊 评估分数：{info['score']}/15**")
                st.markdown(f"**🔄 重写次数：{info['rewrite_count']}次**")
                if info['plan'] and not info['plan'].startswith("❌"):
                    with st.expander("📋 规划摘要"):
                        st.markdown(info['plan'])
                if info['feedback']:
                    with st.expander("💡 评估反馈"):
                        st.markdown(info['feedback'])
                st.markdown("---")
            st.markdown(content)

    if ok_count > 0:
        st.balloons()
            
